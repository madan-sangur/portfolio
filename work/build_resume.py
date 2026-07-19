from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"C:\Users\DELL\Documents\Codex\2026-07-19\h\outputs")
OUT.mkdir(parents=True, exist_ok=True)
DOCX = OUT / "Madan_R_Sangur_Resume.docx"

NAVY = "172033"
BLUE = "2E5EAA"
SLATE = "485363"
LIGHT = "D8DEE8"


def set_run_font(run, name="Aptos", size=10, color=NAVY, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_bottom_border(paragraph, color=LIGHT, size="8"):
    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    pPr.append(borders)


def add_hyperlink(paragraph, text, url, color=BLUE):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_color = OxmlElement("w:color")
    r_color.set(qn("w:val"), color)
    r_pr.append(r_color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "none")
    r_pr.append(u)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_section(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(title.upper())
    set_run_font(run, size=9, color=BLUE, bold=True)
    add_bottom_border(p, color=LIGHT, size="6")
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.19)
    p.paragraph_format.first_line_indent = Inches(-0.14)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.02
    set_run_font(p.add_run(text), size=9.15, color=NAVY)
    return p


def add_project(doc, name, subtitle, bullets):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    set_run_font(p.add_run(name), size=10.2, color=NAVY, bold=True)
    if subtitle:
        set_run_font(p.add_run(f"  |  {subtitle}"), size=9.1, color=SLATE, italic=True)
    for bullet in bullets:
        add_bullet(doc, bullet)


doc = Document()
section = doc.sections[0]
# Named override: resume_compact_margin, needed for one-page recruiter format.
section.top_margin = Inches(0.5)
section.bottom_margin = Inches(0.5)
section.left_margin = Inches(0.58)
section.right_margin = Inches(0.58)
section.header_distance = Inches(0.25)
section.footer_distance = Inches(0.25)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
normal.font.size = Pt(9.2)
normal.font.color.rgb = RGBColor.from_string(NAVY)
normal.paragraph_format.space_after = Pt(2)
normal.paragraph_format.line_spacing = 1.04

# Header pattern: customer_pack adapted as a compact, left-aligned resume masthead.
name = doc.add_paragraph()
name.paragraph_format.space_after = Pt(0)
name.paragraph_format.line_spacing = 0.9
set_run_font(name.add_run("Madan R. Sangur"), size=24, color=NAVY, bold=True)

role = doc.add_paragraph()
role.paragraph_format.space_after = Pt(4)
set_run_font(role.add_run("AI/ML Engineer | B.Tech Artificial Intelligence & Machine Learning Student"), size=10.2, color=BLUE, bold=True)

contact = doc.add_paragraph()
contact.paragraph_format.space_after = Pt(7)
contact.paragraph_format.line_spacing = 1.0
set_run_font(contact.add_run("Haveri, Karnataka, India  |  +91 93803 50086  |  "), size=8.7, color=SLATE)
add_hyperlink(contact, "madansangur@gmail.com", "mailto:madansangur@gmail.com")
set_run_font(contact.add_run("  |  "), size=8.7, color=SLATE)
add_hyperlink(contact, "linkedin.com/in/madan-sangur-075772275", "https://www.linkedin.com/in/madan-sangur-075772275")
set_run_font(contact.add_run("  |  "), size=8.7, color=SLATE)
add_hyperlink(contact, "github.com/madan-sangur", "https://github.com/madan-sangur")
add_bottom_border(contact, color=BLUE, size="12")

add_section(doc, "Professional Summary")
summary = doc.add_paragraph()
summary.paragraph_format.space_after = Pt(2)
summary.paragraph_format.line_spacing = 1.06
set_run_font(summary.add_run(
    "Artificial Intelligence and Machine Learning student with hands-on experience building Python-based ML prototypes and data-driven product concepts. "
    "Experienced with TensorFlow, Keras, Pandas, and NumPy; interested in applying AI to practical, human-centered challenges. Seeking internship and entry-level opportunities in AI/ML and data-focused development."
), size=9.2)

add_section(doc, "Technical Skills")
skills = doc.add_paragraph()
skills.paragraph_format.space_after = Pt(2)
skills.paragraph_format.line_spacing = 1.05
set_run_font(skills.add_run("Languages: "), size=9.1, color=NAVY, bold=True)
set_run_font(skills.add_run("Python, C, Java (OOP), HTML, CSS, JavaScript"), size=9.1)
set_run_font(skills.add_run("    |    ML & Data: "), size=9.1, color=NAVY, bold=True)
set_run_font(skills.add_run("TensorFlow, Keras, Pandas, NumPy, Machine Learning, Deep Learning"), size=9.1)
skills2 = doc.add_paragraph()
skills2.paragraph_format.space_after = Pt(2)
skills2.paragraph_format.line_spacing = 1.05
set_run_font(skills2.add_run("Tools & Platforms: "), size=9.1, color=NAVY, bold=True)
set_run_font(skills2.add_run("Git, GitHub, MySQL (RDBMS), AWS (Basics), Data Visualization"), size=9.1)

add_section(doc, "Projects")
add_project(doc, "ArcLife", "AI-powered life management platform | In development", [
    "Designing a personal growth platform that brings planning, habit tracking, journaling, and goal progress into one focused experience.",
    "Exploring AI-assisted guidance and personalized workflows to help users build sustainable routines."
])
add_project(doc, "NeuroPlay", "AI-driven personalized learning platform | Code Meet 2025 Hackathon", [
    "Designed and prototyped adaptive learning paths based on diagnostic responses and individual aptitude patterns.",
    "Applied Python and ML concepts to categorize aptitude and tailor subject matter to each learner."
])
add_project(doc, "Image Classification Model using CNN", "Python, TensorFlow, Keras", [
    "Developed and optimized a convolutional neural network that achieved approximately 98% validation accuracy."
])
add_project(doc, "Console-based Student Record System", "C, file handling, data structures", [
    "Built a record-management application using C with file handling and core data-structure concepts."
])

add_section(doc, "Education")
edu = doc.add_paragraph()
edu.paragraph_format.space_after = Pt(1)
set_run_font(edu.add_run("B.Tech in Artificial Intelligence & Machine Learning"), size=10.0, color=NAVY, bold=True)
set_run_font(edu.add_run("  |  Srinivas University Institute of Engineering Technology, Mukka"), size=9.1, color=SLATE)
edu2 = doc.add_paragraph()
edu2.paragraph_format.space_after = Pt(2)
set_run_font(edu2.add_run("Expected 2027  |  CGPA: 6.5/10 (after 4th semester)  |  Relevant coursework: Algorithms & Data Structures, Linear Algebra, OOP, Cloud Computing Fundamentals"), size=8.95, color=SLATE)

add_section(doc, "Certifications & Languages")
certs = doc.add_paragraph()
certs.paragraph_format.space_after = Pt(1)
set_run_font(certs.add_run("Certifications: "), size=9.1, color=NAVY, bold=True)
set_run_font(certs.add_run("Data Visualization - IBM Developer Skills Network (2024); RDBMS - IBM Developer Skills Network (2024); Machine Learning - Ethnotech Academic Solutions (2024)"), size=8.9)
langs = doc.add_paragraph()
langs.paragraph_format.space_after = Pt(0)
set_run_font(langs.add_run("Languages: "), size=9.1, color=NAVY, bold=True)
set_run_font(langs.add_run("English (Fluent), Hindi (Conversational), Kannada (Native)"), size=8.9)

doc.core_properties.author = "Madan R. Sangur"
doc.core_properties.title = "Madan R. Sangur - Resume"
doc.core_properties.subject = "AI/ML Engineer Resume"
doc.save(DOCX)
print(DOCX)
