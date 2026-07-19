from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(r"C:\Users\DELL\Documents\Codex\2026-07-19\h\outputs")
OUT.mkdir(parents=True, exist_ok=True)
DOCX = OUT / "Madan_R_Sangur_Resume.docx"
INK = "161B22"
MUTED = "4B5563"


def font(run, size=9.4, bold=None, color=INK, italic=None):
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def hyperlink(paragraph, text, url):
    relation = paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    node = OxmlElement("w:hyperlink")
    node.set(qn("r:id"), relation)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), INK)
    props.append(color)
    run.append(props)
    content = OxmlElement("w:t")
    content.text = text
    run.append(content)
    node.append(run)
    paragraph._p.append(node)


def section(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    font(p.add_run(title.upper()), size=9.3, bold=True)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.14)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.line_spacing = 1.04
    font(p.add_run(text), size=9.15)


def project(doc, name, tech, bullets):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    font(p.add_run(name), size=9.8, bold=True)
    if tech:
        font(p.add_run(f" | {tech}"), size=9.0, color=MUTED)
    for item in bullets:
        bullet(doc, item)


doc = Document()
sec = doc.sections[0]
# Named resume override: compact margins preserve a readable, one-page engineer resume.
sec.top_margin = Inches(0.5)
sec.bottom_margin = Inches(0.5)
sec.left_margin = Inches(0.65)
sec.right_margin = Inches(0.65)
sec.header_distance = Inches(0.25)
sec.footer_distance = Inches(0.25)

normal = doc.styles["Normal"]
normal.font.name = "Aptos"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
normal.font.size = Pt(9.4)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.line_spacing = 1.0

head = doc.add_paragraph()
head.paragraph_format.space_after = Pt(0)
font(head.add_run("Madan R. Sangur"), size=21, bold=True)

role = doc.add_paragraph()
role.paragraph_format.space_after = Pt(4)
font(role.add_run("AI/ML Engineer | B.Tech Artificial Intelligence & Machine Learning Student"), size=9.7, bold=True, color=MUTED)

contact = doc.add_paragraph()
contact.paragraph_format.space_after = Pt(2)
font(contact.add_run("Haveri, Karnataka, India | +91 93803 50086 | "), size=8.7, color=MUTED)
hyperlink(contact, "madansangur@gmail.com", "mailto:madansangur@gmail.com")
font(contact.add_run(" | "), size=8.7, color=MUTED)
hyperlink(contact, "linkedin.com/in/madan-sangur-075772275", "https://www.linkedin.com/in/madan-sangur-075772275")
font(contact.add_run(" | "), size=8.7, color=MUTED)
hyperlink(contact, "github.com/madan-sangur", "https://github.com/madan-sangur")

section(doc, "Summary")
bullet(doc, "Apply Python, TensorFlow, Keras, and data tooling to build machine learning prototypes and practical AI product concepts.")
bullet(doc, "Focus on AI/ML and data-focused internship opportunities where strong fundamentals and hands-on project work can create useful outcomes.")

section(doc, "Projects")
project(doc, "ArcLife", "AI-powered life management platform | In development", [
    "Design a unified experience for planning, habit tracking, journaling, and personal growth.",
    "Explore AI-assisted guidance and personalized workflows that support consistent routines."
])
project(doc, "NeuroPlay", "Python, ML | Code Meet 2025 Hackathon", [
    "Designed and prototyped adaptive learning paths from diagnostic responses and aptitude patterns.",
    "Applied ML concepts to categorize aptitude and tailor subject matter for individual learners."
])
project(doc, "Image Classification Model", "Python, TensorFlow, Keras", [
    "Developed and optimized a convolutional neural network that achieved approximately 98% validation accuracy."
])
project(doc, "Student Record System", "C, file handling, data structures", [
    "Built a console-based record-management application using file handling and core data-structure concepts."
])

section(doc, "Skills")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
font(p.add_run("Programming: "), size=9.1, bold=True)
font(p.add_run("Python, C, Java (OOP), HTML, CSS, JavaScript"), size=9.1)
font(p.add_run("  |  ML & Data: "), size=9.1, bold=True)
font(p.add_run("TensorFlow, Keras, Pandas, NumPy, Machine Learning, Deep Learning"), size=9.1)
p = doc.add_paragraph()
font(p.add_run("Tools: "), size=9.1, bold=True)
font(p.add_run("Git, GitHub, MySQL (RDBMS), AWS (Basics), Data Visualization"), size=9.1)

section(doc, "Education")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
font(p.add_run("B.Tech in Artificial Intelligence & Machine Learning"), size=9.7, bold=True)
font(p.add_run(" | Srinivas University Institute of Engineering Technology, Mukka"), size=9.0, color=MUTED)
p = doc.add_paragraph()
font(p.add_run("Expected 2027 | CGPA: 6.5/10 (after 4th semester) | Coursework: Algorithms & Data Structures, Linear Algebra, OOP, Cloud Computing Fundamentals"), size=8.9, color=MUTED)

section(doc, "Certifications")
p = doc.add_paragraph()
font(p.add_run("Data Visualization, IBM Developer Skills Network (2024) | RDBMS, IBM Developer Skills Network (2024) | Machine Learning, Ethnotech Academic Solutions (2024)"), size=8.75)

doc.core_properties.author = "Madan R. Sangur"
doc.core_properties.title = "Madan R. Sangur Resume"
doc.core_properties.subject = "AI/ML Engineer Resume"
doc.save(DOCX)
print(DOCX)
